"""
DOCX Report Generator for ARIZ sessions.

Uses python-docx to produce editable Word documents.
Structure mirrors the PDF generator:
  Title page -> Problem -> Steps by parts -> Contradictions -> IKR -> Solutions.
"""
import io
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

MODE_LABELS = {
    "express": "Краткий АРИЗ (Экспресс)",
    "full": "Полный АРИЗ-2010",
    "autopilot": "Автопилот",
}

FULL_ARIZ_PARTS = {
    1: "Часть 1: Анализ задачи",
    2: "Часть 2: Анализ ресурсов",
    3: "Часть 3: Определение ОП и ИКР",
    4: "Часть 4: Мобилизация и применение ВПР",
}

EXPRESS_LABELS = {
    "1": "Формулировка задачи",
    "2": "Поверхностное противоречие",
    "3": "Углублённое противоречие",
    "4": "Идеальный конечный результат",
    "5": "Обострённое противоречие",
    "6": "Углубление ОП",
    "7": "Решение",
}

DOMAIN_LABELS = {
    "technical": "Техническая",
    "business": "Бизнес",
    "everyday": "Бытовая",
}

METHOD_LABELS = {
    "principle": "Приём ТРИЗ",
    "standard": "Стандарт",
    "effect": "Эффект",
    "analog": "Аналог",
    "combined": "Комбинированный",
}

CONTRADICTION_TYPE_LABELS = {
    "surface": "Поверхностное (ПП)",
    "deepened": "Углублённое (УП)",
    "sharpened": "Обострённое (ОП)",
}

# Colors
COLOR_PRIMARY = RGBColor(0x1A, 0x23, 0x7E)
COLOR_SECONDARY = RGBColor(0x54, 0x6E, 0x7A)
COLOR_DARK = RGBColor(0x37, 0x47, 0x4F)
COLOR_HEADER_BG = "1565C0"
COLOR_GREEN = RGBColor(0x2E, 0x7D, 0x32)
COLOR_ORANGE = RGBColor(0xEF, 0x6C, 0x00)
COLOR_RED = RGBColor(0xC6, 0x28, 0x28)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, color_hex: str):
    """Set background shading color of a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shading_elem = tc_pr.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): color_hex,
        },
    )
    tc_pr.append(shading_elem)


def _add_run(paragraph, text, bold=False, italic=False, size=None, color=None):
    """Add a formatted run to a paragraph and return it."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    return run


# ---------------------------------------------------------------------------
# Generator
# ---------------------------------------------------------------------------

class DOCXReportGenerator:
    """Generates a DOCX report for a completed ARIZ session.

    Usage::

        generator = DOCXReportGenerator(session)
        docx_bytes = generator.generate()
    """

    def __init__(self, session):
        """Initialize with an ``ARIZSession`` instance."""
        self.session = session
        self.problem = session.problem
        self.doc = Document()
        self._setup_document()

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def generate(self) -> bytes:
        """Generate the DOCX and return raw bytes."""
        self._build_title_page()
        self._build_problem_section()
        self._build_steps_section()
        self._build_contradictions_section()
        self._build_ikr_section()
        self._build_solutions_section()
        self._build_footer()

        buffer = io.BytesIO()
        self.doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        return docx_bytes

    # ------------------------------------------------------------------
    # Setup
    # ------------------------------------------------------------------

    def _setup_document(self):
        """Configure document margins and default styles."""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

        for section in self.doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2)

        self.doc.core_properties.author = "ТРИЗ-Решатель"
        self.doc.core_properties.title = f"АРИЗ-отчёт: {self.problem.title}"

    # ------------------------------------------------------------------
    # Title page
    # ------------------------------------------------------------------

    def _build_title_page(self):
        # Spacing
        for _ in range(5):
            self.doc.add_paragraph("")

        # Title
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p_title, "ТРИЗ-Решатель", bold=True, size=Pt(28), color=COLOR_PRIMARY)

        # Subtitle
        p_sub = self.doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p_sub, "Отчёт по сессии АРИЗ", size=Pt(16), color=COLOR_SECONDARY)

        # Mode
        mode_label = MODE_LABELS.get(self.session.mode, self.session.mode)
        p_mode = self.doc.add_paragraph()
        p_mode.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p_mode, f"Режим: {mode_label}", size=Pt(14), color=COLOR_SECONDARY)

        self.doc.add_paragraph("")
        self.doc.add_paragraph("")

        # Problem title
        p_prob = self.doc.add_paragraph()
        p_prob.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p_prob, "Задача: ", bold=True, size=Pt(13))
        _add_run(p_prob, self.problem.title, size=Pt(13))

        for _ in range(3):
            self.doc.add_paragraph("")

        # Meta info
        date_str = self.session.created_at.strftime("%d.%m.%Y %H:%M")
        self._centered_text(f"Дата создания: {date_str}", color=COLOR_SECONDARY)
        if self.session.completed_at:
            self._centered_text(
                f"Дата завершения: {self.session.completed_at.strftime('%d.%m.%Y %H:%M')}",
                color=COLOR_SECONDARY,
            )
        author_name = self.problem.user.get_full_name() or self.problem.user.username
        self._centered_text(f"Автор: {author_name}", color=COLOR_SECONDARY)

        self.doc.add_page_break()

    # ------------------------------------------------------------------
    # Problem section
    # ------------------------------------------------------------------

    def _build_problem_section(self):
        self._section_heading("1. Описание задачи")

        p = self.doc.add_paragraph()
        _add_run(p, "Название: ", bold=True)
        _add_run(p, self.problem.title)

        domain = DOMAIN_LABELS.get(self.problem.domain, self.problem.domain)
        p2 = self.doc.add_paragraph()
        _add_run(p2, "Область: ", bold=True)
        _add_run(p2, domain)

        if self.problem.original_description:
            p3 = self.doc.add_paragraph()
            _add_run(p3, "Описание:", bold=True)
            self.doc.add_paragraph(self.problem.original_description)

        self.doc.add_paragraph("")

    # ------------------------------------------------------------------
    # Steps section
    # ------------------------------------------------------------------

    def _build_steps_section(self):
        steps = list(self.session.steps.filter(status="completed").order_by("created_at"))
        if not steps:
            return

        self._section_heading("2. Ход решения")

        if self.session.mode == "full":
            self._render_full_steps(steps)
        else:
            self._render_sequential_steps(steps)

    def _render_sequential_steps(self, steps):
        for step in steps:
            label = EXPRESS_LABELS.get(step.step_code, step.step_name)
            self._step_block(step, label)

    def _render_full_steps(self, steps):
        current_part = None
        for step in steps:
            part_num = self._get_part_number(step.step_code)
            if part_num != current_part:
                current_part = part_num
                part_title = FULL_ARIZ_PARTS.get(part_num, f"Часть {part_num}")
                h = self.doc.add_heading(part_title, level=2)
                for run in h.runs:
                    run.font.color.rgb = COLOR_DARK
            self._step_block(step, step.step_name)

    def _step_block(self, step, label: str):
        """Add a single step block."""
        p_title = self.doc.add_paragraph()
        _add_run(p_title, f"Шаг {step.step_code}: {label}", bold=True, size=Pt(11),
                 color=RGBColor(0x28, 0x35, 0x93))

        if step.user_input:
            p = self.doc.add_paragraph()
            _add_run(p, "Ввод пользователя:", bold=True, size=Pt(9), color=COLOR_SECONDARY)
            self.doc.add_paragraph(step.user_input)

        result_text = step.validated_result or step.llm_output
        if result_text:
            p = self.doc.add_paragraph()
            _add_run(p, "Результат:", bold=True, size=Pt(9), color=COLOR_SECONDARY)
            for paragraph in result_text.split("\n\n"):
                paragraph = paragraph.strip()
                if paragraph:
                    self.doc.add_paragraph(paragraph)

        if step.validation_notes:
            p = self.doc.add_paragraph()
            _add_run(p, f"Примечание: {step.validation_notes}", italic=True,
                     size=Pt(9), color=COLOR_SECONDARY)

    # ------------------------------------------------------------------
    # Contradictions section
    # ------------------------------------------------------------------

    def _build_contradictions_section(self):
        contradictions = list(self.session.contradictions.all())
        if not contradictions:
            return

        self._section_heading("3. Противоречия")

        table = self.doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header
        hdr_cells = table.rows[0].cells
        for i, text in enumerate(["Тип", "Формулировка", "Свойство / Анти-свойство"]):
            _set_cell_shading(hdr_cells[i], COLOR_HEADER_BG)
            p = hdr_cells[i].paragraphs[0]
            _add_run(p, text, bold=True, color=COLOR_WHITE, size=Pt(10))

        for c in contradictions:
            row_cells = table.add_row().cells
            row_cells[0].text = CONTRADICTION_TYPE_LABELS.get(c.type, c.type)
            row_cells[1].text = c.formulation[:300]

            props = ""
            if c.property_s or c.anti_property_s:
                props = f"{c.property_s} / {c.anti_property_s}"
            elif c.quality_a or c.quality_b:
                props = f"{c.quality_a} / {c.quality_b}"
            row_cells[2].text = props

        for row in table.rows:
            row.cells[0].width = Cm(3.5)
            row.cells[1].width = Cm(9)
            row.cells[2].width = Cm(4.5)

        self.doc.add_paragraph("")

    # ------------------------------------------------------------------
    # IKR section
    # ------------------------------------------------------------------

    def _build_ikr_section(self):
        ikrs = list(self.session.ikrs.all())
        if not ikrs:
            return

        self._section_heading("4. Идеальный конечный результат (ИКР)")

        for i, ikr in enumerate(ikrs, 1):
            p = self.doc.add_paragraph()
            _add_run(p, f"ИКР-{i}: ", bold=True)
            _add_run(p, ikr.formulation)

            if ikr.strengthened_formulation:
                p2 = self.doc.add_paragraph()
                _add_run(p2, "Усиленная формулировка: ", bold=True)
                _add_run(p2, ikr.strengthened_formulation)

            if ikr.vpr_used:
                vpr_text = ", ".join(str(v) for v in ikr.vpr_used)
                p3 = self.doc.add_paragraph()
                _add_run(p3, "Использованные ВПР: ", bold=True)
                _add_run(p3, vpr_text)

            self.doc.add_paragraph("")

    # ------------------------------------------------------------------
    # Solutions section
    # ------------------------------------------------------------------

    def _build_solutions_section(self):
        solutions = list(self.session.solutions.all())
        if not solutions:
            return

        self._section_heading("5. Решения")

        for i, sol in enumerate(solutions, 1):
            method_label = METHOD_LABELS.get(sol.method_used, sol.method_used)

            h = self.doc.add_heading(f"Решение {i}: {sol.title}", level=2)
            for run in h.runs:
                run.font.color.rgb = COLOR_DARK

            p_method = self.doc.add_paragraph()
            _add_run(p_method, "Метод: ", bold=True)
            _add_run(p_method, method_label)

            self.doc.add_paragraph(sol.description)

            # Scores table
            score_table = self.doc.add_table(rows=3, cols=3)
            score_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            score_table.style = "Table Grid"

            # Header row
            for j, txt in enumerate(["Показатель", "Оценка", "Уровень"]):
                _set_cell_shading(score_table.cell(0, j), "283593")
                p = score_table.cell(0, j).paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_run(p, txt, bold=True, color=COLOR_WHITE, size=Pt(10))

            # Novelty row
            score_table.cell(1, 0).text = "Новизна"
            p_n = score_table.cell(1, 1).paragraphs[0]
            p_n.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p_n, f"{sol.novelty_score}/10", bold=True,
                     color=self._score_color(sol.novelty_score))
            score_table.cell(1, 2).text = self._score_label(sol.novelty_score)

            # Feasibility row
            score_table.cell(2, 0).text = "Реализуемость"
            p_f = score_table.cell(2, 1).paragraphs[0]
            p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p_f, f"{sol.feasibility_score}/10", bold=True,
                     color=self._score_color(sol.feasibility_score))
            score_table.cell(2, 2).text = self._score_label(sol.feasibility_score)

            for row in score_table.rows:
                row.cells[0].width = Cm(4)
                row.cells[1].width = Cm(3)
                row.cells[2].width = Cm(5)

            self.doc.add_paragraph("")

    # ------------------------------------------------------------------
    # Footer
    # ------------------------------------------------------------------

    def _build_footer(self):
        self.doc.add_paragraph("")
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(
            p,
            f"Сгенерировано ТРИЗ-Решателем {datetime.now().strftime('%d.%m.%Y %H:%M')}",
            size=Pt(8),
            color=COLOR_SECONDARY,
        )

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _section_heading(self, text: str):
        """Add a section heading (Heading 1) with primary color."""
        h = self.doc.add_heading(text, level=1)
        for run in h.runs:
            run.font.color.rgb = COLOR_PRIMARY
            run.font.size = Pt(16)

    def _centered_text(self, text: str, color=None):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, text, color=color, size=Pt(11))

    @staticmethod
    def _get_part_number(step_code: str) -> int:
        try:
            return int(step_code.split(".")[0])
        except (ValueError, IndexError):
            return 1

    @staticmethod
    def _score_color(score: int) -> RGBColor:
        if score >= 7:
            return COLOR_GREEN
        if score >= 4:
            return COLOR_ORANGE
        return COLOR_RED

    @staticmethod
    def _score_label(score: int) -> str:
        if score >= 8:
            return "Отлично"
        if score >= 6:
            return "Хорошо"
        if score >= 4:
            return "Средне"
        if score >= 2:
            return "Ниже среднего"
        return "Низко"
